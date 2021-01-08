from functools import lru_cache
import docx
import time
import collections
import os
# inputDoc = docx.Document('Tolstoj.docx')
filesFolder = 'FilesToAnalise'
testText = ['Красивые на яблоки\n на лежали на столе', 'Десять на негритят на сосали на сушку']


def get_all_files(folder):
    allFiles = os.listdir(folder)
    return allFiles

def get_text(doc):
    text = []
    doc = docx.Document(filesFolder+"/"+doc)
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    return text


def split_words(textArray):
    words = []
    for par in textArray:
        parWords = par.split(' ')
        for word in parWords:
            words.append(str(word))
    return words


def split_words_logic(textArray):
    if len(textArray) == 1:
        return split_words(textArray[0])
    else:
        multiWords = []
        for array in textArray:
            for word in split_words(array):
                multiWords.append(word)
        return multiWords


def files_analise_logic(filesNames):
    if len(filesNames) == 1:
        return get_text(filesNames[0])
    else:
        multiText = []
        for file in filesNames:
            multiText.append(get_text(file))
        return multiText


char_to_replace = {'(': '',
                   ')': '',
                   '{': '',
                   '}': '',
                   '[': '',
                   ']': '',
                   ',': '',
                   '.': '',
                   '+': '',
                   '|': '',
                   '>': '',
                   '<': '',
                   '*': '',
                   '\\': '',
                   '/': ''
                   }


def filter_symbols(word):
    filteredWord = word
    for key, value in char_to_replace.items():
        filteredWord = filteredWord.replace(key, value)
    return filteredWord.rstrip()


def filter_del_nums(word):
    if not [s for s in word if s in '0123456789']:
        return word


def filter_only_letters(word):
    if word.isalpha():
        return word


def filter_to_lower(word):
    lowWord = str(word).lower()
    return lowWord


def word_filter(wordArray):
    res = []
    for word in wordArray:
        word = filter_symbols(word)
        word = filter_del_nums(word)
        word = filter_to_lower(word)
        filtered_final = word
        if filtered_final != 'none' and filtered_final != '' and filtered_final is not None:
            res.append(word)
    return res


def unique_words(wordArray):
    unique = set(wordArray)
    return unique


def word_counter(wordArray):
    c = collections.Counter(wordArray)
    return c.most_common()


def counted_words_printer(words, top, letters):
    goodWords = []
    i = 0
    while len(goodWords) < top:
        if i >= len(words):
            break
        if len(words[i][0]) > letters:
            goodWords.append(words[i])
        i += 1
    for item in goodWords:
        print(f"{item[0]} - {item[1]} раз")


def text_analise():
    # -----------------------------------------
    print('Парсинг текста из docx файла...')
    startParse = time.time()
    parsedText = files_analise_logic(get_all_files(filesFolder))
    endParse = time.time()
    print(f"Выполнено за: {endParse - startParse:.2f} секунд")
    # -----------------------------------------
    print('Разбивка текста на слова...')
    startSplit = time.time()
    splittedWords = split_words_logic(parsedText)
    endSplit = time.time()
    print(f"Выполнено за: {endSplit - startSplit:.2f} секунд")
    # -----------------------------------------
    print('Фильтруем полученные данные...')
    startFilter = time.time()
    filteredWords = word_filter(splittedWords)
    endFilter = time.time()
    print(f"Выполнено за: {endFilter - startFilter:.2f} секунд")
    # -----------------------------------------
    print('Отсеиваем уникальные символьные ряды...')
    startUniq = time.time()
    uniqWords = unique_words(filteredWords)
    endUniq = time.time()
    print(f"Выполнено за: {endUniq - startUniq:.2f} секунд")
    # -----------------------------------------
    print('Считаем количество повторений каждого слова...')
    startCount = time.time()
    countWords = word_counter(filteredWords)
    endCount = time.time()
    print(f"Выполнено за: {endCount - startCount:.2f} секунд")
    # -----------------------------------------
    print("----------------Stats----------------")
    print(f"Количество слов: {len(splittedWords)} || с фильтром {len(filteredWords)}")
    print(f"Количество уникальных слов: {len(uniqWords)}")
    print(f"----------------Top words----------------")
    counted_words_printer(countWords, 10, 2)

text_analise()

